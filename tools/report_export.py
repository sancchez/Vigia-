"""Exportación de reportes markdown a PDF/DOCX — item transversal de HANDOFF.md
("reportes descargables (PDF/DOCX)").

Ambos reportes que ya existen en el proyecto (`agents/reporteria.py` para un
scan técnico, `agents/cumplimiento.py` para el reporte de cumplimiento
normativo) generan su salida como texto markdown vía el LLM. Este módulo NO
vuelve a generar contenido: solo convierte ese markdown ya producido a bytes
descargables, en el servidor, para que la descarga sea igual sin importar el
navegador del cliente.

Elección de librerías (ver HANDOFF.md / handoff de esta sesión): `fpdf2` y
`python-docx` ya estaban instaladas en el entorno del proyecto y son puras en
Python — sin binarios de sistema (wkhtmltopdf, GTK/Pango de weasyprint) que
son dolorosos de instalar en Windows. A cambio, el parser de markdown de
abajo es deliberadamente simple: encabezados (`#`/`##`/`###`), listas con
`-`/`*`, negrita `**texto**` y párrafos. Tablas markdown (`| col | col |`)
se degradan a texto plano línea por línea — una limitación conocida, no un
renderer de markdown completo.
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from fpdf import FPDF

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _strip_bold(text: str) -> str:
    return _BOLD_RE.sub(r"\1", text)


_TYPOGRAFIA_A_ASCII = {
    "—": "-",  # em dash — Claude lo usa constantemente en sus reportes
    "–": "-",  # en dash
    "‘": "'",
    "’": "'",
    "“": '"',
    "”": '"',
    "…": "...",
    "•": "-",  # bullet suelto dentro de texto (no al inicio de línea)
}


def _latin1(text: str) -> str:
    """Las fuentes core de fpdf2 (helvetica) solo cubren latin-1.

    El acento español (á, é, í, ó, ú, ñ, ¿, ¡) SÍ está en latin-1, pero la
    puntuación "tipográfica" que Claude usa todo el tiempo en sus reportes
    (— em dash, comillas curvas, …) NO — sin este reemplazo previo, cada
    em dash del reporte real se convertía en un "?" suelto (bug real
    encontrado leyendo el PDF generado contra un reporte de verdad, no una
    suposición). Emoji/símbolos verdaderamente raros sí se reemplazan con
    "?" en vez de reventar la generación completa del PDF.
    """
    for original, ascii_equivalente in _TYPOGRAFIA_A_ASCII.items():
        text = text.replace(original, ascii_equivalente)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def markdown_to_pdf_bytes(markdown_text: str, titulo: str = "Reporte") -> bytes:
    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def _linea(texto: str, alto: float) -> None:
        # `new_x="LMARGIN", new_y="NEXT"` es obligatorio aquí: el default de
        # fpdf2 (`new_x=RIGHT`) deja el cursor pegado al margen derecho
        # después de cada `multi_cell`, así que la SIGUIENTE línea revienta
        # con "Not enough horizontal space to render a single character" —
        # bug real encontrado probando esto contra un reporte de verdad, no
        # una suposición.
        pdf.multi_cell(0, alto, _latin1(texto), new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("helvetica", "B", 16)
    _linea(titulo, 10)
    pdf.ln(2)

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            pdf.ln(3)
            continue
        if stripped in ("---", "***"):
            pdf.ln(2)
            continue
        if line.startswith("### "):
            pdf.set_font("helvetica", "B", 12)
            _linea(_strip_bold(line[4:]), 8)
        elif line.startswith("## "):
            pdf.set_font("helvetica", "B", 13)
            _linea(_strip_bold(line[3:]), 8)
        elif line.startswith("# "):
            pdf.set_font("helvetica", "B", 15)
            _linea(_strip_bold(line[2:]), 9)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            pdf.set_font("helvetica", "", 10.5)
            _linea(f"  - {_strip_bold(stripped[2:])}", 6)
        else:
            pdf.set_font("helvetica", "", 10.5)
            _linea(_strip_bold(line), 6)

    return bytes(pdf.output())


def _add_runs(paragraph, text: str) -> None:
    """Agrega texto a un párrafo de python-docx, respetando **negrita** como runs separados."""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx_bytes(markdown_text: str, titulo: str = "Reporte") -> bytes:
    doc = Document()
    doc.add_heading(titulo, level=0)

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped in ("---", "***"):
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _add_runs(doc.add_paragraph(style="List Bullet"), stripped[2:])
        else:
            _add_runs(doc.add_paragraph(), line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
